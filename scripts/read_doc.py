import sys
import os
from docx import Document
from pypdf import PdfReader

def read_docx(file_path):
    doc = Document(file_path)
    output = []
    
    # Process elements in order if possible, or just paragraphs and tables
    # For simplicity and coverage, we'll iterate through block-level elements
    for element in doc.element.body:
        if element.tag.endswith('p'):
            # Find the paragraph object
            for p in doc.paragraphs:
                if p._element == element:
                    if p.text.strip():
                        output.append(p.text)
                    break
        elif element.tag.endswith('tbl'):
            # Find the table object
            for t in doc.tables:
                if t._element == element:
                    output.append("\n--- Table ---")
                    for row in t.rows:
                        row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                        # Remove duplicate adjacent cells caused by merged cells
                        unique_row_data = []
                        for cell in row_data:
                            if not unique_row_data or unique_row_data[-1] != cell:
                                unique_row_data.append(cell)
                        output.append(" | ".join(unique_row_data))
                    output.append("-------------\n")
                    break
    
    # Fallback to simple paragraphs if structure extraction was incomplete
    if not output:
        for para in doc.paragraphs:
            if para.text.strip():
                output.append(para.text)
                
    return "\n".join(output)

def read_pdf(file_path):
    reader = PdfReader(file_path)
    output = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        output.append(f"--- Page {i + 1} ---")
        output.append(text)
    return "\n".join(output)

def main():
    if len(sys.argv) < 2:
        print("Usage: python read_doc.py <path_to_file>", file=sys.stderr)
        sys.exit(1)
        
    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(f"Error: File not found: {file_path}", file=sys.stderr)
        sys.exit(1)
        
    _, ext = os.path.splitext(file_path.lower())
    
    try:
        if ext == '.docx':
            print(read_docx(file_path))
        elif ext == '.pdf':
            print(read_pdf(file_path))
        else:
            print(f"Error: Unsupported file format '{ext}'. Only .docx and .pdf are supported.", file=sys.stderr)
            sys.exit(1)
    except Exception as e:
        print(f"Error reading file: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    main()
